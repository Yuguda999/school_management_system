"""
Utility functions for converting documents between formats.
"""
import re
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_table_borders(table):
    """Add borders to a table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Create border elements
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    
    tblPr.append(tblBorders)


def markdown_to_docx(markdown_text: str, output_path: str, title: Optional[str] = None) -> str:
    """
    Convert markdown text to a DOCX file.
    
    Args:
        markdown_text: The markdown content to convert
        output_path: Path where the DOCX file should be saved
        title: Optional title to add at the top of the document
        
    Returns:
        The path to the created DOCX file
    """
    doc = Document()
    
    # Add title if provided
    if title:
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add spacing
    
    # Split content into lines
    lines = markdown_text.split('\n')
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle tables
        if '|' in line and not in_table:
            # Start of a table
            in_table = True
            table_data = [line]
            i += 1
            continue
        elif in_table:
            if '|' in line:
                table_data.append(line)
                i += 1
                continue
            else:
                # End of table, process it
                process_markdown_table(doc, table_data)
                table_data = []
                in_table = False
                doc.add_paragraph()  # Add spacing after table
                continue
        
        # Handle headings
        if line.startswith('# '):
            para = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            para = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            para = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            para = doc.add_heading(line[5:].strip(), level=4)
        
        # Handle bullet lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, text)
        
        # Handle numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            para = doc.add_paragraph(style='List Number')
            add_formatted_text(para, text)
        
        # Handle horizontal rules
        elif line.strip() in ['---', '***', '___']:
            para = doc.add_paragraph()
            para.add_run('_' * 50)
        
        # Handle empty lines
        elif not line.strip():
            doc.add_paragraph()
        
        # Handle regular paragraphs
        else:
            para = doc.add_paragraph()
            add_formatted_text(para, line)
        
        i += 1
    
    # Process any remaining table
    if in_table and table_data:
        process_markdown_table(doc, table_data)
    
    # Save the document
    doc.save(output_path)
    return output_path


def process_markdown_table(doc: Document, table_lines: list):
    """Process markdown table and add it to the document"""
    if not table_lines:
        return
    
    # Parse table rows
    rows = []
    for line in table_lines:
        # Skip separator lines (e.g., |---|---|)
        if re.match(r'^\|[\s\-:]+\|$', line.strip()):
            continue
        
        # Split by | and clean up
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last elements from split
        cells = [c for c in cells if c]
        
        if cells:
            rows.append(cells)
    
    if not rows:
        return
    
    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Add borders
    add_table_borders(table)
    
    # Fill table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                # First row is header
                if i == 0:
                    cell_para = cell.paragraphs[0]
                    run = cell_para.add_run(cell_text)
                    run.font.bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    cell.text = cell_text


def add_formatted_text(paragraph, text: str):
    """Add text with markdown formatting (bold, italic) to a paragraph"""
    # Handle bold and italic
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold with **
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        # Bold with __
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        # Italic with *
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        # Italic with _
        elif part.startswith('_') and part.endswith('_') and not part.startswith('__'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        # Code with `
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        # Regular text
        else:
            paragraph.add_run(part)

